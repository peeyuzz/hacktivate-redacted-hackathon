import fitz
import re
import pytesseract
from PIL import Image
import os
import json
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
from dotenv import load_dotenv
import cv2
import numpy as np
from faceDetection import FaceDetector
import platform
from scipy.ndimage import rotate
from deskew import determine_skew
import whisper
from pydub import AudioSegment
import docx
from docx.shared import RGBColor
import openpyxl
import pandas as pd

load_dotenv()

API_KEY = os.environ["API_KEY"]
genai.configure(api_key=API_KEY)

# Set tesseract command path based on OS
if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract'
elif platform.system() == "Linux":
    pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

class Redactor:
    # Regex patterns (unchanged)
    EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@(?:[A-Za-z0-9-]+\.)+[A-Z|a-z]{2,}\b'
    PHONE_PATTERN = r'\b(?:\+?\d{1,4}[\s-]?)?(?:(?:\d{3}[\s-]?){2}\d{4}|\d{10})\b'
    AADHAR_PATTERN = r'\b[2-9]\d{3}\s?-?\d{4}\s?-?\d{4}\b'
    PASSPORT_PATTERN = r'\b[A-PR-WY][1-9]\d{7}\b'
    PAN_PATTERN = r'\b[A-Z]{5}[0-9]{4}[A-Z]\b'
    BANK_ACC_PATTERN = r'\b\d{9,18}\b'
    CREDIT_CARD_PATTERN = r'\b(?:4\d{3}|5[1-5]\d{2}|6011|65\d{2}|3[47]\d|30[012345]\d)\d{11}\b'
    DRIVING_LICENSE_PATTERN = r'\b[A-Z]{2}\d{13}\b'
    VEHICLE_REGISTRATION_PATTERN = r'\b[A-Z]{2}\d{1,2}[A-Z]{1,3}\d{1,4}\b'

    def __init__(self, path, plan_type="free", special_instructions=None, level = ["low", "medium","high"]):
        self.path = path
        self.special_instructions = special_instructions
        self.plan_type = plan_type
        model_path = os.path.join("model", "public", "ultra-lightweight-face-detection-rfb-320", "FP16", "ultra-lightweight-face-detection-rfb-320.xml")
        self.face_detector = FaceDetector(model=model_path)
        self.zoom_factor = 2  # Consolidated zoom factor
        self.level = level

    def preprocess_image(self, image):
        # Convert to grayscale if needed
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image

        # Detect and correct skew
        angle = determine_skew(gray)
        rotated = rotate(gray, angle, reshape=False, mode='constant', cval=255, order=3)

        # Apply sharpening
        kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
        sharpened = cv2.filter2D(rotated, -1, kernel)

        return sharpened

    def get_text_with_ocr(self, image):
        preprocessed = self.preprocess_image(image)
        ocr_data = pytesseract.image_to_data(preprocessed, output_type=pytesseract.Output.DICT)
        
        full_text = " ".join(ocr_data['text'])
        blocks = []
        for i in range(len(ocr_data['text'])):
            if ocr_data['text'][i].strip():
                x, y, w, h = ocr_data['left'][i], ocr_data['top'][i], ocr_data['width'][i], ocr_data['height'][i]
                blocks.append([x, y, x+w, y+h, ocr_data['text'][i]])
        
        return full_text, blocks

    def get_sensitive_data(self, text):
        if self.plan_type == 'pro':
            with open('prompt.txt', 'r', encoding='utf-8') as file:
                prompt = file.read()
            if self.special_instructions:
                prompt += f"\n SPECIAL INSTRUCTIONS FROM THE USER, THESE HAVE THE HIGHEST PRIORITY: {self.special_instructions}"
            
            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(
                f"<DOCUMENT>\n{text}\n</DOCUMENT>\n\n{prompt}",
                safety_settings={
                    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
                })
            
            try:
                response_text = response.text
                match = re.search(r"```json\n(.*?)\n```", response_text, re.DOTALL)
                if match:
                    json_str = match.group(1).strip()
                    return json.loads(json_str)
                else:
                    print("No JSON content found")
                    return []
            except json.JSONDecodeError as e:
                print(f"JSON decoding error: {e}")
                return []
        else:
            patterns = [
                re.compile(pattern, re.IGNORECASE) for pattern in [
                    self.EMAIL_PATTERN, self.PHONE_PATTERN, self.AADHAR_PATTERN,
                    self.PASSPORT_PATTERN, self.PAN_PATTERN, self.BANK_ACC_PATTERN,
                    self.CREDIT_CARD_PATTERN, self.DRIVING_LICENSE_PATTERN,
                    self.VEHICLE_REGISTRATION_PATTERN
                ]
            ]
            
            sensitive_data = []
            for pattern in patterns:
                matches = pattern.finditer(text)
                for match in matches:
                    sensitive_data.append({"text": match.group(), "reason": "regex match", "level": "high"})
            return sensitive_data

    def find_text_locations(self, text_blocks, sensitive_data):
        locations = []
        full_text = " ".join([block[4].lower() for block in text_blocks])
        
        for data in sensitive_data:
            if data['level'] in self.level:
                data_text = data['text'].lower()
                start = full_text.find(data_text)
                if start != -1:
                    end = start + len(data_text)
                    word_start = full_text[:start].count(' ')
                    word_end = word_start + data_text.count(' ') + 1
                    
                    for block in text_blocks[word_start:word_end]:
                        x0, y0, x1, y1 = block[:4]
                    #    locations.append(fitz.Rect(x0, y0, x1, y1))
                        locations.append(fitz.Rect(x0, y0, x1, y1))
            else:
                continue
        return locations

    def get_face_and_qr_locations(self, image):
        faces, _ = self.face_detector.inference(image)
        
        # face_locations = [fitz.Rect(face) for face in faces]
        face_locations = []
        for face in faces:
            x0, y0, x1, y1 = face
            face_locations.append(fitz.Rect(x0, y0, x1, y1))

        qr_detector = cv2.QRCodeDetector()
        _, _, points, _ = qr_detector.detectAndDecodeMulti(image)
        qr_locations = []
        if points is not None:
            for point in points:
                x0, y0, x1, y1 = point[0][0], point[0][1], point[2][0], point[2][1]
                qr_locations.append(fitz.Rect(x0, y0, x1, y1))

        return face_locations + qr_locations

    def redact_pdf(self):
        doc = fitz.open(self.path)
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(self.zoom_factor, self.zoom_factor))
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            image_np = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

            text, text_blocks = self.get_text_with_ocr(image_np)
            sensitive_data = self.get_sensitive_data(text)
            
            locations = self.find_text_locations(text_blocks, sensitive_data)
            locations.extend(self.get_face_and_qr_locations(image_np))

            for location in locations:
                page.add_redact_annot(location / self.zoom_factor, fill=(0, 0, 0))
            page.apply_redactions()

        output_path = os.path.splitext(self.path)[0] + '_redacted.pdf'
        doc.save(output_path)
        print(f"Successfully redacted and saved as {output_path}")
        return locations
    
    def redact_txt(self):
        with open(self.path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        sensitive_data = self.get_sensitive_data(content)
        redacted_content = content

        for data in sensitive_data:
            if data['level'] in self.level:
                redacted_content = redacted_content.replace(data['text'], '[REDACTED]')

        output_path = os.path.splitext(self.path)[0] + '_redacted.txt'
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(redacted_content)
        
        print(f"Successfully redacted and saved as {output_path}")

    def redact_docx(self):
        doc = docx.Document(self.path)
        content = "\n".join([para.text for para in doc.paragraphs])
        
        sensitive_data = self.get_sensitive_data(content)
        print(sensitive_data)
        for para in doc.paragraphs:
            for data in sensitive_data:
                if data['level'] in self.level and data['text'] in para.text:
                    para.text = para.text.replace(data['text'], '[REDACTED]')

        output_path = os.path.splitext(self.path)[0] + '_redacted.docx'
        doc.save(output_path)
        print(f"Successfully redacted and saved as {output_path}")

    def redact_xlsx(self):
        df = pd.read_excel(self.path)
        content = df.to_string()
        
        sensitive_data = self.get_sensitive_data(content)

        for data in sensitive_data:
            if data['level'] in self.level:
                df = df.replace(data['text'], '[REDACTED]', regex=True)

        output_path = os.path.splitext(self.path)[0] + '_redacted.xlsx'
        df.to_excel(output_path, index=False)
        print(f"Successfully redacted and saved as {output_path}")

    def redact_image(self):
        image = cv2.imread(self.path)
        height, width = image.shape[:2]
        zoomed_image = cv2.resize(image, (width * self.zoom_factor, height * self.zoom_factor), interpolation=cv2.INTER_LINEAR)

        text, text_blocks = self.get_text_with_ocr(zoomed_image)
        sensitive_data = self.get_sensitive_data(text)
        
        locations = self.find_text_locations(text_blocks, sensitive_data)
        locations.extend(self.get_face_and_qr_locations(zoomed_image))

        for location in locations:
            x_min, y_min, x_max, y_max = np.array(location / self.zoom_factor, dtype=np.int32)
            cv2.rectangle(image, (x_min, y_min), (x_max, y_max), (0, 0, 0), -1)

        output_path = os.path.splitext(self.path)[0] + '_redacted.jpg'
        cv2.imwrite(output_path, image)
        print(f"Successfully redacted image and saved as {output_path}")
        return locations
    
    def redact_video(self):
        video = cv2.VideoCapture(self.path)
        fps = video.get(cv2.CAP_PROP_FPS)
        width = int(video.get(cv2.CAP_PROP_FRAME_WIDTH))
        height = int(video.get(cv2.CAP_PROP_FRAME_HEIGHT))

        output_path = os.path.splitext(self.path)[0] + '_redacted.mp4'
        fourcc = cv2.VideoWriter_fourcc(*'mp4v')
        out = cv2.VideoWriter(output_path, fourcc, fps, (width, height))

        while True:
            ret, frame = video.read()
            if not ret:
                break
            
            # Zoom the frame
            zoomed_frame = cv2.resize(frame, (width * self.zoom_factor, height * self.zoom_factor), interpolation=cv2.INTER_LINEAR)
            
            # Get text and sensitive data (commented out for performance, uncomment if needed)
            # text, text_blocks = self.get_text_with_ocr(zoomed_frame)
            # sensitive_data = self.get_sensitive_data(text)
            # text_locations = self.find_text_locations(text_blocks, sensitive_data)
            
            # Get face locations
            face_locations = self.get_face_and_qr_locations(zoomed_frame)
            
            # Combine all locations
            # locations = text_locations + face_locations
            locations = face_locations  # If using only face detection
            
            # Apply redaction
            for location in locations:
                x_min, y_min, x_max, y_max = np.array(location / self.zoom_factor, dtype=np.int32)
                cv2.rectangle(frame, (x_min, y_min), (x_max, y_max), (0, 0, 0), -1)
                
            out.write(frame)

        video.release()
        out.release()
        print(f"Successfully redacted video and saved as {output_path}")

    def transcribe_audio_with_whisper(self):
        model = whisper.load_model("tiny")
        transcript = model.transcribe(audio=self.path, word_timestamps=True)
        return transcript
    
    def find_audio_locations(self, transcript, sensitive_data):
        locations = []
        for data in sensitive_data:
            for segment in transcript['segments']:
                if data['text'].lower() in segment['text'].lower():
                    start_time = segment['start']
                    end_time = segment['end']
                    
                    # Fine-tune start and end times
                    for word in segment['words']:
                        if word['word'].lower().strip() in data['text'].lower():
                            start_time = word['start']
                            break
                    
                    for word in reversed(segment['words']):
                        if word['word'].lower().strip() in data['text'].lower():
                            end_time = word['end']
                            break
                    
                    locations.append((start_time, end_time))
        return locations
    
    def redact_audio(self):
        audio = AudioSegment.from_file(self.path)
        transcript = self.transcribe_audio_with_whisper()
        text = ' '.join([segment['text'] for segment in transcript['segments']])
        sensitive_data = self.get_sensitive_data(text)
        locations = self.find_audio_locations(transcript, sensitive_data)
        beep = AudioSegment.from_wav("beep.wav")
        
        redacted_audio = AudioSegment.empty()
        last_end = 0

        for start_time, end_time in sorted(locations):
            start_ms = int(start_time * 1000)
            end_ms = int(end_time * 1000)
            redacted_audio += audio[last_end:start_ms]
            redacted_audio += beep[:end_ms - start_ms]
            last_end = end_ms
        redacted_audio += audio[last_end:]
        
        output_path = os.path.splitext(self.path)[0] + '_redacted.wav'
        redacted_audio.export(output_path, format="wav")
        print(f"Successfully redacted audio and saved as {output_path}")

    def redact(self):
        file_extension = os.path.splitext(self.path)[1].lower()
        redaction_methods = {
            ".pdf": self.redact_pdf,
            ".mp3": self.redact_audio,
            ".wav": self.redact_audio,
            ".jpg": self.redact_image,
            ".jpeg": self.redact_image,
            ".png": self.redact_image,
            ".mp4": self.redact_video,
            ".avi": self.redact_video,
            ".mov": self.redact_video,
            ".txt": self.redact_txt,
            ".docx": self.redact_docx,
            ".xlsx": self.redact_xlsx
        }
        
        redact_method = redaction_methods.get(file_extension)
        if redact_method:
            return redact_method()
        else:
            print(f"Unsupported file type: {file_extension}")
            return None

# path = r"C:\Users\admin\Documents\RMSI itnern\Intern Information Form Filled (1).docx"
# redactor = Redactor(path, plan_type="pro", level=["low"])
# redactor.redact()